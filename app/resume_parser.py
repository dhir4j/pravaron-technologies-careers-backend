from __future__ import annotations

import html
import re
import zlib
import zipfile
from io import BytesIO

from docx import Document
from pypdf import PdfReader
from pathlib import Path

MAX_PDF_FALLBACK_RAW_BYTES = 10 * 1024 * 1024
MAX_PDF_FALLBACK_STREAMS = 120
MAX_PDF_FALLBACK_INFLATED_BYTES = 2 * 1024 * 1024



def repair_spaced_char_words(value: str) -> str:
    repaired_lines: list[str] = []
    for line in value.splitlines():
        segments = re.split(r" {2,}", line)
        repaired_segments: list[str] = []
        for segment in segments:
            tokens = segment.split(" ")
            meaningful = [token for token in tokens if token]
            if len(meaningful) >= 3:
                singleish = [token for token in meaningful if len(token) == 1 or token in {"+", "-", "/", ".", "@", "|", ":", ","}]
                if len(singleish) / len(meaningful) >= 0.72:
                    repaired_segments.append("".join(meaningful))
                    continue
            repaired_segments.append(segment)
        repaired_lines.append(" ".join(part for part in repaired_segments if part.strip()))
    return "\n".join(repaired_lines)

def normalize_resume_text(value: str) -> str:
    value = value.replace("\x00", " ")
    value = "".join(ch if ch.isprintable() or ch in "\n\t" else " " for ch in value)
    value = repair_spaced_char_words(value)
    value = re.sub(r"[ \t\r\f\v]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return "\n".join(line.strip() for line in value.splitlines() if line.strip()).strip()


def extract_docx_text(raw: bytes) -> str:
    try:
        document = Document(BytesIO(raw))
        chunks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                chunks.append(" | ".join(cell.text for cell in row.cells))
        text = normalize_resume_text("\n".join(chunks))
        if text:
            return text
    except Exception:
        pass
    with zipfile.ZipFile(BytesIO(raw)) as archive:
        names = [name for name in archive.namelist() if name.startswith("word/") and name.endswith(".xml")]
        ordered = ["word/document.xml"] + [name for name in names if name != "word/document.xml"]
        chunks: list[str] = []
        for name in ordered:
            if name not in archive.namelist():
                continue
            xml = archive.read(name).decode("utf-8", errors="ignore")
            xml = re.sub(r"</w:p>", "\n", xml)
            xml = re.sub(r"</w:tr>", "\n", xml)
            xml = re.sub(r"<[^>]+>", " ", xml)
            chunks.append(html.unescape(xml))
        return normalize_resume_text("\n".join(chunks))


def _pdf_unescape(value: str) -> str:
    value = re.sub(r"\\([nrtbf])", lambda m: {"n": "\n", "r": "\n", "t": "\t", "b": "", "f": ""}[m.group(1)], value)
    value = re.sub(r"\\([()\\])", r"\1", value)
    value = re.sub(r"\\([0-7]{1,3})", lambda m: chr(int(m.group(1), 8)), value)
    return value


def _extract_pdf_strings(data: bytes) -> str:
    text = data.decode("latin-1", errors="ignore")
    strings = [_pdf_unescape(match) for match in re.findall(r"\((?:\\.|[^\\)])*\)", text)]
    hex_strings = []
    for value in re.findall(r"<([0-9A-Fa-f\s]{8,})>", text):
        compact = re.sub(r"\s+", "", value)
        if len(compact) % 2:
            compact = compact[:-1]
        try:
            decoded = bytes.fromhex(compact)
        except ValueError:
            continue
        for encoding in ("utf-16-be", "utf-8", "latin-1"):
            try:
                candidate = decoded.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
            if candidate:
                hex_strings.append(candidate)
                break
    return normalize_resume_text("\n".join(strings + hex_strings))


def _merge_unique_text_chunks(chunks: list[str]) -> str:
    seen: set[str] = set()
    lines: list[str] = []
    for chunk in chunks:
        for line in normalize_resume_text(chunk).splitlines():
            key = re.sub(r"\s+", " ", line).strip().lower()
            if not key or key in seen:
                continue
            seen.add(key)
            lines.append(line)
    return normalize_resume_text("\n".join(lines))


def _safe_zlib_decompress(data: bytes) -> bytes | None:
    decompressor = zlib.decompressobj()
    try:
        inflated = decompressor.decompress(data, MAX_PDF_FALLBACK_INFLATED_BYTES + 1)
    except zlib.error:
        return None
    if len(inflated) > MAX_PDF_FALLBACK_INFLATED_BYTES:
        return None
    return inflated


def _extract_pdf_fallback_chunks(raw: bytes) -> list[str]:
    if len(raw) > MAX_PDF_FALLBACK_RAW_BYTES:
        return []
    chunks = [_extract_pdf_strings(raw)]
    for index, match in enumerate(re.finditer(rb"stream\r?\n(.*?)\r?\nendstream", raw, flags=re.S)):
        if index >= MAX_PDF_FALLBACK_STREAMS:
            break
        stream = match.group(1).strip(b"\r\n")
        for candidate in (stream, stream.strip()):
            inflated = _safe_zlib_decompress(candidate)
            if inflated is None:
                continue
            extracted = _extract_pdf_strings(inflated)
            if extracted:
                chunks.append(extracted)
            break
    return chunks


def extract_pdf_text(raw: bytes) -> str:
    chunks: list[str] = []
    page_count = 0
    extracted_pages = 0
    try:
        reader = PdfReader(BytesIO(raw))
        if getattr(reader, "is_encrypted", False):
            try:
                reader.decrypt("")
            except Exception:
                pass
        page_count = len(reader.pages)
        for page in reader.pages:
            try:
                page_text = page.extract_text() or ""
            except Exception:
                continue
            if page_text.strip():
                extracted_pages += 1
                chunks.append(page_text)
    except Exception:
        pass

    if not chunks or (page_count and extracted_pages < page_count):
        chunks.extend(_extract_pdf_fallback_chunks(raw))
    return _merge_unique_text_chunks([chunk for chunk in chunks if chunk])


def extract_binary_doc_text(raw: bytes) -> str:
    text = raw.decode("latin-1", errors="ignore")
    words = re.findall(r"[A-Za-z0-9@._+:/#,&()\-]{2,}", text)
    return normalize_resume_text(" ".join(words))


def extract_resume_text(filename: str, raw: bytes, content_type: str | None = None) -> tuple[str, str, str | None]:
    ext = Path(filename or "").suffix.lower().lstrip(".")
    try:
        if ext == "docx" or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = extract_docx_text(raw)
        elif ext == "pdf" or content_type == "application/pdf":
            text = extract_pdf_text(raw)
        elif ext == "doc":
            text = extract_binary_doc_text(raw)
        else:
            text = extract_binary_doc_text(raw)
        if not text:
            return "", "empty", "No readable text could be extracted from this resume."
        return text, "extracted", None
    except Exception as exc:
        return "", "failed", str(exc)
